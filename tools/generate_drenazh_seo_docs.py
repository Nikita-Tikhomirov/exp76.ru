# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Pt


ROOT = Path(r"C:\Users\user\Desktop\exp76.ru\seo-content\drenazh-uchastka")
SUBFOLDERS = [
    "01-core",
    "02-subservices",
    "03-geo-backlog",
    "04-problem-backlog",
    "99-admin-checklists",
]


def ensure_dirs() -> None:
    for sub in SUBFOLDERS:
        (ROOT / sub).mkdir(parents=True, exist_ok=True)


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(value)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def write_service_doc(path: Path, data: dict) -> None:
    doc = Document()
    set_default_style(doc)
    doc.add_heading(data["name"], level=1)

    add_kv(doc, "URL", data["url"])
    add_kv(doc, "Slug", data["slug"])
    add_kv(doc, "Шаблон", data["template"])
    add_kv(doc, "Intent", data["intent"])
    add_kv(doc, "Основной кластер", data["main_query"])

    doc.add_heading("SEO-мета", level=2)
    add_kv(doc, "Title", data["title"])
    add_kv(doc, "H1", data["h1"])
    add_kv(doc, "Description", data["description"])

    doc.add_heading("Вторичные запросы (3-6)", level=2)
    add_bullets(doc, data["secondary"])

    if data["template"] == "category-87.php":
        doc.add_heading("Контент в редакторе (the_content)", level=2)
        for paragraph in data["content"]:
            doc.add_paragraph(paragraph)
        doc.add_heading("ACF-поля cat87_*", level=2)
    else:
        doc.add_heading("Блоки под newservicepost (ns87_*)", level=2)

    for key, value in data["acf"].items():
        add_kv(doc, key, value)

    doc.add_heading("Перелинковка", level=2)
    add_bullets(doc, data["links"])
    doc.add_heading("CTA", level=2)
    add_kv(doc, "Основной CTA", data["cta"])
    doc.save(path)


def write_geo_backlog() -> None:
    cities = ["yaroslavl", "rybinsk", "uglich", "tutaev", "pereslavl"]
    for i, city in enumerate(cities, start=1):
        doc = Document()
        set_default_style(doc)
        doc.add_heading(f"Гео-страница дренажа: {city}", level=1)
        add_kv(doc, "URL", f"/{city}/drenazh-uchastka/")
        add_kv(doc, "Статус", "Backlog / ТЗ на подготовку контента")
        doc.add_heading("Что подготовить перед написанием", level=2)
        add_bullets(
            doc,
            [
                "Собрать локальные кейсы и фото работ по городу.",
                "Проверить локальные ценовые ориентиры.",
                "Добавить FAQ с учетом местных условий и грунтов.",
                "Подготовить локальные офферы и сроки выезда.",
            ],
        )
        doc.add_heading("Шаблон SEO-мета", level=2)
        add_kv(doc, "Title", f"Дренаж участка в {city} — монтаж под ключ")
        add_kv(doc, "H1", f"Дренаж участка в {city} под ключ")
        add_kv(
            doc,
            "Description",
            f"Дренаж участка в {city} и районе: расчет, монтаж, гарантия, выезд инженера.",
        )
        doc.save(ROOT / "03-geo-backlog" / f"{i:02d}-{city}-drenazh-uchastka-backlog.docx")


def write_problem_backlog() -> None:
    problems = [
        ("bolotistyj-uchastok", "Болотистый участок"),
        ("gryaznyj-uchastok", "Грязный участок"),
        ("posle-dozhdya-voda", "Вода после дождя"),
    ]
    for i, (slug, name) in enumerate(problems, start=1):
        doc = Document()
        set_default_style(doc)
        doc.add_heading(f"Проблемная страница: {name}", level=1)
        add_kv(doc, "URL", f"/drenazh-uchastka/{slug}/")
        add_kv(doc, "Статус", "Backlog / ТЗ")
        doc.add_heading("Структура страницы", level=2)
        add_bullets(
            doc,
            [
                "Описание проблемы и рисков для фундамента.",
                "Рабочее решение и тип дренажа.",
                "Пример сметы и сроки.",
                "FAQ и призыв к заявке.",
            ],
        )
        doc.save(ROOT / "04-problem-backlog" / f"{i:02d}-{slug}-backlog.docx")


def write_checklists() -> None:
    doc = Document()
    set_default_style(doc)
    doc.add_heading("Чек-лист загрузки в WordPress", level=1)
    add_bullets(
        doc,
        [
            "Проверить URL и slug без дублей.",
            "Заполнить Title / H1 / Description.",
            "Заполнить обязательные ACF-поля.",
            "Проверить Hero, FAQ, цены, CTA.",
            "Проверить перелинковку между страницами.",
            "Проверить мобильное отображение и кириллицу.",
        ],
    )
    doc.save(ROOT / "99-admin-checklists" / "01-checklist-wp-load.docx")

    doc = Document()
    set_default_style(doc)
    doc.add_heading("Матрица ACF-полей для дренажа", level=1)
    doc.add_heading("category-87.php", level=2)
    add_bullets(
        doc,
        [
            "cat87_hero_*",
            "cat87_offer_title, cat87_trust_items",
            "cat87_prices_* и cat87_estimate_*",
            "cat87_faq_*",
            "the_content для SEO-текста",
        ],
    )
    doc.add_heading("newservicepost.php", level=2)
    add_bullets(
        doc,
        [
            "ns87_hero_*",
            "ns87_problem_*",
            "ns87_solution_*",
            "ns87_prices_* и ns87_estimate_*",
            "ns87_faq_*",
        ],
    )
    add_kv(
        doc,
        "Важно",
        "В newservicepost.php сейчас фиксированный контекст category_87; лучше перейти на контекст текущего поста.",
    )
    doc.save(ROOT / "99-admin-checklists" / "02-acf-field-matrix.docx")


def main() -> None:
    ensure_dirs()
    pages = [
        {
            "folder": "01-core",
            "filename": "01-drenazh-uchastka-glavnaya.docx",
            "name": "Дренаж участка — главная страница услуги",
            "url": "/drenazh-uchastka/",
            "slug": "drenazh-uchastka",
            "template": "category-87.php",
            "intent": "Коммерческий: заказ дренажа под ключ в Ярославской области",
            "main_query": "дренаж участка",
            "title": "Дренаж участка под ключ в Ярославской области — цена и монтаж",
            "h1": "Дренаж участка под ключ в Ярославской области",
            "description": "Проектирование и монтаж дренажа участка в Ярославле и области. Расчет стоимости за 1 день.",
            "secondary": [
                "дренаж участка под ключ",
                "устройство дренажа на участке",
                "монтаж дренажа на участке",
                "дренаж участка стоимость",
                "системы дренажа участка",
            ],
            "content": [
                "Если после дождя на участке стоит вода, разрушается отмостка и сыреет фундамент, проблему решает правильно рассчитанная дренажная система.",
                "Мы проектируем и монтируем дренаж в Ярославле и по Ярославской области с учетом типа грунта, уровня грунтовых вод и рельефа.",
                "Итог работ — сухой участок, защищенный фундамент и прогнозируемые расходы на обслуживание.",
            ],
            "acf": {
                "cat87_hero_title": "Дренаж участка под ключ в Ярославской области",
                "cat87_hero_subtitle": "Проектируем и монтируем рабочие дренажные системы. Расчет стоимости за 1 день.",
                "cat87_hero_btn_primary_text": "Рассчитать стоимость",
                "cat87_hero_btn_primary_url": "#calc",
                "cat87_hero_btn_secondary_text": "Получить консультацию",
                "cat87_hero_btn_secondary_url": "#popup",
                "cat87_offer_title": "Что вы получаете при заказе дренажа под ключ",
                "cat87_prices_title": "Стоимость дренажа участка",
                "cat87_faq_title": "Частые вопросы по дренажу",
            },
            "links": [
                "/drenazh-uchastka/cena/",
                "/drenazh-uchastka/vysokie-gruntovye-vody/",
                "/drenazh-uchastka/glinistaya-pochva/",
                "/drenazh-uchastka/s-uklonom/",
            ],
            "cta": "Оставьте заявку и получите расчет дренажа участка в течение 1 рабочего дня.",
        },
        {
            "folder": "01-core",
            "filename": "02-drenazh-uchastka-cena.docx",
            "name": "Дренаж участка — страница цены",
            "url": "/drenazh-uchastka/cena/",
            "slug": "cena",
            "template": "newservicepost.php",
            "intent": "Коммерческий: расчет и сравнение стоимости",
            "main_query": "дренаж участка цена",
            "title": "Цена дренажа участка под ключ в Ярославле и области",
            "h1": "Стоимость дренажа участка: цены и примеры смет",
            "description": "Актуальные цены на дренаж участка под ключ: за метр и за объект. Смета с учетом грунта и УГВ.",
            "secondary": [
                "дренаж участка под ключ цена",
                "сколько стоит дренаж участка",
                "стоимость дренажа участка",
                "дренаж участка цена за метр",
            ],
            "acf": {
                "ns87_hero_title": "Стоимость дренажа участка под ключ",
                "ns87_hero_subtitle": "Прозрачные цены: за метр, за этап и за готовый объект.",
                "ns87_problem_title": "Почему итоговая цена отличается",
                "ns87_solution_title": "Как мы считаем стоимость",
                "ns87_prices_title": "Цены на дренажные работы",
                "ns87_faq_title": "Вопросы по цене",
            },
            "links": [
                "/drenazh-uchastka/",
                "/drenazh-uchastka/10-sotok/",
                "/drenazh-uchastka/6-sotok/",
                "/drenazh-uchastka/s-uklonom/",
            ],
            "cta": "Оставьте параметры участка и получите предварительную смету в день обращения.",
        },
    ]

    subservices = [
        ("vysokie-gruntovye-vody", "Дренаж при высоких грунтовых водах"),
        ("glinistaya-pochva", "Дренаж глинистого участка"),
        ("vokrug-doma", "Дренаж вокруг дома"),
        ("glubinnyy", "Глубинный дренаж"),
        ("poverhnostnyy", "Поверхностный дренаж"),
        ("10-sotok", "Дренаж участка 10 соток"),
        ("6-sotok", "Дренаж участка 6 соток"),
        ("s-uklonom", "Дренаж участка с уклоном"),
    ]
    for slug, name in subservices:
        pages.append(
            {
                "folder": "02-subservices",
                "filename": f"{slug}.docx",
                "name": name,
                "url": f"/drenazh-uchastka/{slug}/",
                "slug": slug,
                "template": "newservicepost.php",
                "intent": "Коммерческий: услуга под ключ",
                "main_query": name.lower(),
                "title": f"{name} в Ярославле и Ярославской области",
                "h1": f"{name} под ключ",
                "description": f"{name}. Выезд на участок, расчет сметы, монтаж под ключ и гарантия.",
                "secondary": [
                    "дренаж участка цена",
                    "дренаж участка под ключ",
                    "монтаж дренажа на участке",
                    "стоимость дренажа участка",
                ],
                "acf": {
                    "ns87_hero_title": f"{name} под ключ",
                    "ns87_hero_subtitle": "Подбираем схему под задачу участка и монтируем с гарантией.",
                    "ns87_problem_title": "Типовые проблемы участка",
                    "ns87_solution_title": "Как решаем задачу",
                    "ns87_prices_title": "Стоимость работ",
                    "ns87_faq_title": "Вопросы по услуге",
                },
                "links": [
                    "/drenazh-uchastka/",
                    "/drenazh-uchastka/cena/",
                    "/drenazh-uchastka/glubinnyy/",
                    "/drenazh-uchastka/poverhnostnyy/",
                ],
                "cta": "Оставьте заявку, чтобы получить расчет и рабочее решение под ваш участок.",
            }
        )

    for item in pages:
        write_service_doc(ROOT / item["folder"] / item["filename"], item)

    write_geo_backlog()
    write_problem_backlog()
    write_checklists()


if __name__ == "__main__":
    main()
